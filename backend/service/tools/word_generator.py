"""
Word Document Generator Tool
Word 문서 자동 생성 도구
"""

from typing import Dict, Any, List, Optional
from pathlib import Path
from datetime import datetime
import json
import logging
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


logger = logging.getLogger(__name__)


class WordGenerator:
    """Word 문서 생성기"""

    def __init__(self, template_dir: Optional[Path] = None):
        """
        초기화

        Args:
            template_dir: 템플릿 파일 디렉토리
        """
        self.template_dir = template_dir or Path("./backend/service/templates")
        self.output_dir = Path("./generated_documents")
        self.output_dir.mkdir(exist_ok=True)

    def create_document_from_template(
        self,
        template_name: str,
        data: Dict[str, Any],
        output_filename: Optional[str] = None
    ) -> str:
        """
        템플릿을 기반으로 Word 문서 생성

        Args:
            template_name: 템플릿 이름 (예: 'product_seminar_application')
            data: 문서에 채울 데이터
            output_filename: 출력 파일명 (미지정시 자동 생성)

        Returns:
            생성된 파일 경로
        """
        try:
            # 템플릿 로드
            template = self._load_template(template_name)

            # 문서 생성
            doc = Document()

            # 문서 설정
            self._set_document_properties(doc, template.get("properties", {}))

            # 테이블 생성 (Word 문서의 주요 구조)
            if template.get("table_structure"):
                table = self._create_table_from_template(
                    doc,
                    template["table_structure"],
                    data
                )

            # 파일 저장
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"{template_name}_{timestamp}.docx"

            output_path = self.output_dir / output_filename
            doc.save(str(output_path))

            logger.info(f"Word document created: {output_path}")
            return str(output_path)

        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            raise

    def _load_template(self, template_name: str) -> Dict[str, Any]:
        """템플릿 JSON 파일 로드"""
        template_path = self.template_dir / f"{template_name}.json"

        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        with open(template_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    def _set_document_properties(self, doc: Document, properties: Dict[str, Any]):
        """문서 속성 설정"""
        # 페이지 설정
        sections = doc.sections
        for section in sections:
            # A4 사이즈 설정
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)

            # 여백 설정
            section.top_margin = Cm(properties.get("top_margin", 2.0))
            section.bottom_margin = Cm(properties.get("bottom_margin", 2.0))
            section.left_margin = Cm(properties.get("left_margin", 2.0))
            section.right_margin = Cm(properties.get("right_margin", 2.0))

    def _create_table_from_template(
        self,
        doc: Document,
        template: Dict[str, Any],
        data: Dict[str, Any]
    ) -> Any:
        """템플릿 기반 테이블 생성"""

        # 테이블 생성
        rows = template.get("rows", 1)
        cols = template.get("cols", 1)
        table = doc.add_table(rows=rows, cols=cols)

        # 테이블 스타일 설정
        table.style = template.get("style", "Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 각 행 처리
        row_templates = template.get("row_templates", [])

        for row_idx, row_template in enumerate(row_templates):
            if row_idx >= len(table.rows):
                # 필요시 행 추가
                table.add_row()

            row = table.rows[row_idx]
            self._process_table_row(row, row_template, data)

        # 동적 행 추가 (예: 참석자 목록)
        dynamic_sections = template.get("dynamic_sections", [])
        for section in dynamic_sections:
            self._add_dynamic_rows(table, section, data)

        return table

    def _process_table_row(
        self,
        row: Any,
        row_template: Dict[str, Any],
        data: Dict[str, Any]
    ):
        """테이블 행 처리"""

        cells = row_template.get("cells", [])

        for cell_idx, cell_template in enumerate(cells):
            if cell_idx >= len(row.cells):
                continue

            cell = row.cells[cell_idx]

            # 셀 내용 설정
            content = cell_template.get("content", "")
            field_key = cell_template.get("field", None)

            if field_key and field_key in data:
                # 데이터에서 값 가져오기
                content = str(data[field_key])
            elif "{" in content and "}" in content:
                # 템플릿 변수 치환
                content = self._replace_variables(content, data)

            # 텍스트 설정
            cell.text = content

            # 스타일 설정
            if cell_template.get("bold", False):
                self._set_cell_bold(cell)

            if cell_template.get("align"):
                self._set_cell_alignment(cell, cell_template["align"])

            # 셀 병합 처리
            if cell_template.get("merge_with"):
                merge_idx = cell_template["merge_with"]
                if merge_idx < len(row.cells):
                    cell.merge(row.cells[merge_idx])

    def _add_dynamic_rows(
        self,
        table: Any,
        section: Dict[str, Any],
        data: Dict[str, Any]
    ):
        """동적 행 추가 (예: 참석자 목록)"""

        data_key = section.get("data_key")
        if not data_key or data_key not in data:
            return

        items = data[data_key]
        if not isinstance(items, list):
            return

        # 각 아이템에 대해 행 추가
        for item in items:
            row = table.add_row()

            # 컬럼 매핑
            column_mapping = section.get("columns", [])
            for col_idx, col_config in enumerate(column_mapping):
                if col_idx >= len(row.cells):
                    continue

                cell = row.cells[col_idx]
                field = col_config.get("field")

                if field and isinstance(item, dict):
                    cell.text = str(item.get(field, ""))
                elif isinstance(item, str):
                    if col_idx == 0:  # 첫 번째 컬럼에만 값 설정
                        cell.text = item

    def _replace_variables(self, template_str: str, data: Dict[str, Any]) -> str:
        """템플릿 변수 치환 (예: {date} -> 실제 날짜)"""
        result = template_str

        # 기본 변수들
        variables = {
            "{date}": datetime.now().strftime("%Y-%m-%d"),
            "{time}": datetime.now().strftime("%H:%M"),
            "{year}": str(datetime.now().year),
            "{month}": str(datetime.now().month),
            "{day}": str(datetime.now().day),
        }

        # 데이터에서 변수 추가
        for key, value in data.items():
            variables[f"{{{key}}}"] = str(value)

        # 치환
        for var, value in variables.items():
            result = result.replace(var, value)

        return result

    def _set_cell_bold(self, cell):
        """셀 텍스트를 굵게 설정"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    def _set_cell_alignment(self, cell, alignment: str):
        """셀 정렬 설정"""
        align_map = {
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }

        if alignment in align_map:
            for paragraph in cell.paragraphs:
                paragraph.alignment = align_map[alignment]

    def create_product_seminar_application(self, data: Dict[str, Any]) -> str:
        """
        제품설명회 신청서 생성

        Args:
            data: 신청서 데이터
                - seminar_type: 단일/복수
                - pm_attendance: PM 참석 여부
                - date: 일시
                - location: 장소
                - product_name: 제품명
                - expected_attendees: 참석 예정 인원
                - purpose: 시행 목적
                - main_content: 주요 내용
                - staff_list: 직원 목록 [{name, team, ...}, ...]
                - hcp_list: 보건의료전문가 목록 [{name, hospital, ...}, ...]

        Returns:
            생성된 파일 경로
        """
        # 파일명 생성 (콜론을 언더스코어로 변경)
        date_str = data.get('date', datetime.now().strftime('%Y%m%d_%H%M'))
        date_str = date_str.replace(':', '_').replace(' ', '_').replace('-', '')
        filename = f"제품설명회_신청서_{date_str}.docx"

        return self.create_document_from_template(
            template_name="product_seminar_application",
            data=data,
            output_filename=filename
        )

    def create_product_seminar_report(self, data: Dict[str, Any]) -> str:
        """
        제품설명회 결과보고서 생성

        Args:
            data: 결과보고서 데이터
                - seminar_type: 단일/복수
                - pm_attendance: PM 참석 여부
                - date: 일시
                - location: 장소
                - product_name: 제품명
                - actual_attendees: 실제 참석 인원
                - result: 시행 결과
                - main_content: 주요 내용
                - payment_details: 지급 내역
                - budget_usage: 예산 사용 내역
                - staff_list: 직원 목록
                - hcp_list: 보건의료전문가 목록

        Returns:
            생성된 파일 경로
        """
        # 파일명 생성 (콜론을 언더스코어로 변경)
        date_str = data.get('date', datetime.now().strftime('%Y%m%d_%H%M'))
        date_str = date_str.replace(':', '_').replace(' ', '_').replace('-', '')
        filename = f"제품설명회_결과보고서_{date_str}.docx"

        return self.create_document_from_template(
            template_name="product_seminar_report",
            data=data,
            output_filename=filename
        )