"""
Word 문서 구조 분석 도구
"""

import sys
import os
import io
from pathlib import Path

# Windows 인코딩 설정
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(
        sys.stdout.buffer,
        encoding='utf-8',
        errors='replace'
    )

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("python-docx 라이브러리가 필요합니다.")
    print("설치: pip install python-docx")
    sys.exit(1)


def analyze_word_document(doc_path):
    """Word 문서 구조 분석"""

    print(f"\n{'='*80}")
    print(f"📄 Document: {Path(doc_path).name}")
    print(f"{'='*80}")

    try:
        doc = Document(doc_path)

        # 1. 문서 속성
        print("\n📋 Document Properties:")
        print(f"  - Core properties available: {hasattr(doc, 'core_properties')}")
        if hasattr(doc, 'core_properties'):
            core = doc.core_properties
            print(f"    - Title: {core.title}")
            print(f"    - Author: {core.author}")
            print(f"    - Created: {core.created}")
            print(f"    - Modified: {core.modified}")
            print(f"    - Subject: {core.subject}")

        # 2. 섹션 정보
        print(f"\n📑 Sections: {len(doc.sections)}")
        for i, section in enumerate(doc.sections):
            print(f"  Section {i+1}:")
            print(f"    - Page width: {section.page_width}")
            print(f"    - Page height: {section.page_height}")
            print(f"    - Orientation: {section.orientation}")

        # 3. 단락 분석
        print(f"\n📝 Paragraphs: {len(doc.paragraphs)}")
        style_count = {}
        alignment_count = {}

        for para in doc.paragraphs:
            # 스타일 수집
            style = para.style.name if para.style else "None"
            style_count[style] = style_count.get(style, 0) + 1

            # 정렬 수집
            alignment = str(para.alignment) if para.alignment else "None"
            alignment_count[alignment] = alignment_count.get(alignment, 0) + 1

        print(f"\n  Styles used:")
        for style, count in sorted(style_count.items(), key=lambda x: -x[1])[:10]:
            print(f"    - {style}: {count}")

        print(f"\n  Alignments:")
        for align, count in alignment_count.items():
            print(f"    - {align}: {count}")

        # 4. 텍스트 내용 샘플
        print(f"\n📖 Content Sample (first 10 non-empty paragraphs):")
        non_empty_count = 0
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() and non_empty_count < 10:
                text = para.text.strip()[:100]
                if len(para.text.strip()) > 100:
                    text += "..."
                print(f"  {non_empty_count+1}. [{para.style.name if para.style else 'No style'}] {text}")
                non_empty_count += 1

        # 5. 테이블 분석
        print(f"\n📊 Tables: {len(doc.tables)}")
        for i, table in enumerate(doc.tables):
            print(f"  Table {i+1}:")
            print(f"    - Rows: {len(table.rows)}")
            print(f"    - Columns: {len(table.columns)}")
            if len(table.rows) > 0 and len(table.columns) > 0:
                # 첫 번째 행 내용 (헤더일 가능성)
                first_row = table.rows[0]
                headers = []
                for cell in first_row.cells:
                    text = cell.text.strip()[:30]
                    if text:
                        headers.append(text)
                if headers:
                    print(f"    - First row: {headers}")

        # 6. 이미지/도형
        # python-docx로는 inline shapes 정보만 제한적으로 접근 가능
        inline_shapes_count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if hasattr(run, '_element'):
                    if run._element.xpath('.//w:drawing'):
                        inline_shapes_count += 1

        print(f"\n🖼️  Inline shapes/images detected: {inline_shapes_count}")

        # 7. 스타일 정보
        print(f"\n🎨 Available Styles:")
        style_types = {
            WD_STYLE_TYPE.PARAGRAPH: "Paragraph",
            WD_STYLE_TYPE.CHARACTER: "Character",
            WD_STYLE_TYPE.TABLE: "Table",
            WD_STYLE_TYPE.LIST: "List"
        }

        style_summary = {}
        for style in doc.styles:
            style_type = style_types.get(style.type, "Other")
            style_summary[style_type] = style_summary.get(style_type, 0) + 1

        for style_type, count in style_summary.items():
            print(f"  - {style_type} styles: {count}")

        # 8. 특수 요소 찾기
        print(f"\n🔍 Special Elements:")

        # 폼 필드 표시 찾기 (대괄호 또는 밑줄)
        form_fields = []
        for para in doc.paragraphs:
            text = para.text
            # 대괄호 패턴 [___] 또는 [필드명]
            import re
            brackets = re.findall(r'\[([^\]]+)\]', text)
            if brackets:
                form_fields.extend(brackets)

        if form_fields:
            print(f"  Form fields detected: {len(set(form_fields))}")
            print(f"    Examples: {list(set(form_fields))[:5]}")

        # 페이지 나누기 확인
        page_breaks = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if hasattr(run, '_element'):
                    if run._element.xpath('.//w:br[@w:type="page"]'):
                        page_breaks += 1

        print(f"  Page breaks: {page_breaks}")

        return True

    except Exception as e:
        print(f"❌ Error analyzing document: {e}")
        return False


def main():
    """메인 함수"""

    doc_dir = r"C:\kdy\Projects\narutalk_upgrade\beta_v0033\database\storage\documents"

    print("="*80)
    print("📚 Word Document Structure Analysis")
    print("="*80)

    # Word 파일 목록
    doc_files = [
        "제품설명회 시행 결과보고서.docx",
        "제품설명회 시행 신청서.docx"
    ]

    for doc_file in doc_files:
        doc_path = os.path.join(doc_dir, doc_file)
        if os.path.exists(doc_path):
            analyze_word_document(doc_path)
        else:
            print(f"\n⚠️ File not found: {doc_file}")

    print("\n" + "="*80)
    print("✅ Analysis Complete")
    print("="*80)


if __name__ == "__main__":
    main()