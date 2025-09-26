"""
Word 문서 상세 구조 분석 도구
"""

import sys
import os
import io
from pathlib import Path

# Windows 인코딩 설정
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(
        sys.stdout.buffer,
        encoding='utf-8',
        errors='replace'
    )

try:
    from docx import Document
except ImportError:
    print("python-docx 라이브러리가 필요합니다.")
    print("설치: pip install python-docx")
    sys.exit(1)


def analyze_table_structure(doc_path):
    """Word 문서의 테이블 구조 상세 분석"""

    print(f"\n{'='*80}")
    print(f"📄 Detailed Analysis: {Path(doc_path).name}")
    print(f"{'='*80}")

    try:
        doc = Document(doc_path)

        # 테이블 상세 분석
        for table_idx, table in enumerate(doc.tables):
            print(f"\n📊 Table {table_idx + 1}: {len(table.rows)} rows x {len(table.columns)} columns")
            print("-" * 60)

            # 각 행의 내용 출력
            for row_idx, row in enumerate(table.rows):
                row_content = []
                merged_info = []

                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    # 빈 셀이 아니면 내용 추가
                    if text:
                        # 50자로 제한
                        if len(text) > 50:
                            text = text[:47] + "..."
                        row_content.append(f"Cell {cell_idx}: {text}")

                    # 병합 셀 확인
                    if cell_idx > 0 and row.cells[cell_idx] == row.cells[cell_idx - 1]:
                        if cell_idx not in merged_info:
                            merged_info.append(f"Cells {cell_idx-1}-{cell_idx} merged")

                # 행 출력
                if row_content or row_idx < 5:  # 처음 5행은 비어있어도 출력
                    print(f"\n  Row {row_idx + 1}:")
                    if row_content:
                        for content in row_content:
                            print(f"    {content}")
                    else:
                        print(f"    [Empty row]")

                    if merged_info:
                        print(f"    Merge info: {', '.join(merged_info)}")

                # 10행 이상이면 중략
                if row_idx >= 10 and row_idx < len(table.rows) - 2:
                    if row_idx == 11:
                        print(f"\n  ... ({len(table.rows) - 13} more rows) ...")
                    continue

            # 테이블 구조 분석
            print(f"\n  📋 Table Structure Analysis:")

            # 병합된 셀 패턴 확인
            has_merged_cells = False
            for row in table.rows:
                for i in range(1, len(row.cells)):
                    if row.cells[i] == row.cells[i-1]:
                        has_merged_cells = True
                        break
                if has_merged_cells:
                    break

            print(f"    - Has merged cells: {has_merged_cells}")

            # 첫 번째 행이 헤더인지 확인
            if len(table.rows) > 0:
                first_row_texts = [cell.text.strip() for cell in table.rows[0].cells]
                if all(first_row_texts):
                    print(f"    - Potential header row: Yes")
                    print(f"      Headers: {first_row_texts}")
                else:
                    print(f"    - Potential header row: No")

            # 셀 스타일 분석 (첫 번째 행)
            if len(table.rows) > 0:
                print(f"\n  🎨 Cell Styles (First Row):")
                for cell_idx, cell in enumerate(table.rows[0].cells[:4]):  # 최대 4개 셀만
                    if cell.paragraphs:
                        para = cell.paragraphs[0]
                        if para.runs:
                            run = para.runs[0]
                            print(f"    Cell {cell_idx + 1}:")
                            print(f"      - Bold: {run.bold}")
                            print(f"      - Font size: {run.font.size}")
                            print(f"      - Font name: {run.font.name}")

        # 문서 내 필드 패턴 찾기
        print(f"\n🔍 Field Patterns in Document:")

        field_patterns = {}
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    # 빈 칸 또는 입력 필드 패턴
                    import re

                    # 패턴 1: 라벨: 값
                    label_value = re.findall(r'([가-힣\w\s]+):\s*([가-힣\w\s]*)', text)
                    for label, value in label_value:
                        if label.strip():
                            field_patterns[label.strip()] = value.strip() if value.strip() else "[입력 필드]"

                    # 패턴 2: 대괄호 [___]
                    brackets = re.findall(r'\[([^\]]*)\]', text)
                    for bracket in brackets:
                        if bracket and not bracket.strip().startswith('_'):
                            field_patterns[f"[{bracket}]"] = "[대괄호 필드]"

        if field_patterns:
            print(f"  Found {len(field_patterns)} field patterns:")
            for field, value in list(field_patterns.items())[:20]:  # 최대 20개
                print(f"    - {field}: {value}")

        return True

    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False


def extract_template_structure(doc_path):
    """템플릿 구조 추출 - JSON 형식으로"""

    try:
        doc = Document(doc_path)
        doc_name = Path(doc_path).stem

        template_structure = {
            "document_name": doc_name,
            "sections": [],
            "fields": [],
            "tables": []
        }

        # 테이블 구조 추출
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                "table_index": table_idx,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "data": []
            }

            # 각 행의 데이터 추출
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data["data"].append(row_data)

            template_structure["tables"].append(table_data)

            # 필드 추출
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and not text.isspace():
                        # 필드명과 같은 패턴 추출
                        import re
                        if re.search(r'[가-힣]+\s*:', text) or re.search(r'\[.*\]', text):
                            template_structure["fields"].append(text)

        # JSON으로 저장
        import json
        output_file = f"{doc_name}_structure.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(template_structure, f, ensure_ascii=False, indent=2)

        print(f"\n✅ Template structure saved to: {output_file}")

        return template_structure

    except Exception as e:
        print(f"❌ Error extracting template: {e}")
        return None


def main():
    """메인 함수"""

    doc_dir = r"C:\kdy\Projects\narutalk_upgrade\beta_v0033\database\storage\documents"

    print("="*80)
    print("📚 Word Document Detailed Structure Analysis")
    print("="*80)

    # Word 파일 목록
    doc_files = [
        "제품설명회 시행 신청서.docx",  # 신청서 먼저
        "제품설명회 시행 결과보고서.docx"  # 결과보고서
    ]

    for doc_file in doc_files:
        doc_path = os.path.join(doc_dir, doc_file)
        if os.path.exists(doc_path):
            analyze_table_structure(doc_path)
            extract_template_structure(doc_path)
        else:
            print(f"\n⚠️ File not found: {doc_file}")

    print("\n" + "="*80)
    print("✅ Analysis Complete")
    print("="*80)


if __name__ == "__main__":
    main()